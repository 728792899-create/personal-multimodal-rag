from __future__ import annotations

import argparse
import json
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
sys.path.insert(0, str(BACKEND))

from app.services.document_processor import DocumentProcessor  # noqa: E402
from app.services.document_registry import DocumentRegistry  # noqa: E402
from app.services.graph_store import NativeGraphStore  # noqa: E402
from app.services.rag_engine import RagEngine  # noqa: E402
from app.services.retriever import HybridRetriever  # noqa: E402


METRIC_LABELS = {
    "recall_at_5": "Recall@5",
    "mrr": "MRR",
    "citation_accuracy": "引用准确率",
    "citation_coverage": "可验证事实引用覆盖率",
    "valid_citation_rate": "有效引用率",
    "refusal_f1": "拒答 F1",
    "refusal_recall": "无答案拒答召回率",
    "answerable_misrefusal_complement": "可回答问题非误拒率",
    "answer_acceptance_accuracy": "可回答接受率",
    "modality_recall_at_10": "多模态 Recall@10",
    "table_cell_accuracy": "表格单元准确率",
    "caption_alignment": "Caption 对齐率",
    "formula_accuracy": "公式准确率",
    "graph_path_precision": "Graph 路径精度",
    "graph_evidence_coverage": "Graph 证据覆盖率",
    "multihop_complete_evidence_at_10": "多跳完整证据链@10",
}


def load_jsonl(path: Path) -> list[dict]:
    cases: list[dict] = []
    seen: set[str] = set()
    for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        if not line.strip():
            continue
        case = json.loads(line)
        case_id = str(case.get("id") or f"line-{line_number}")
        if case_id in seen:
            raise ValueError(f"评测 Case ID 重复：{case_id}")
        if not str(case.get("question", "")).strip():
            raise ValueError(f"评测 Case {case_id} 缺少问题")
        seen.add(case_id)
        case["id"] = case_id
        cases.append(case)
    if not cases:
        raise ValueError(f"未在 {path} 中找到评测 Case")
    return cases


def build_offline_engine(documents_dir: Path) -> RagEngine:
    processor = DocumentProcessor()
    registry = DocumentRegistry(":memory:")
    graph_store = NativeGraphStore(registry)
    retriever = HybridRetriever(graph_store=graph_store)
    paths = sorted(documents_dir.glob("*.md"))
    if not paths:
        raise ValueError(f"未在 {documents_dir} 中找到 Markdown fixture")
    for path in paths:
        document = processor.parse_file(path)
        document.metadata["knowledge_base_id"] = "operations" if path.name.startswith(("04-", "05-")) else "default"
        retriever.add_document(document, processor.split(document))
        if document.metadata["knowledge_base_id"] == "default":
            registry.save_document(document)
            graph_store.build_document(document)
    for spec_path in sorted(documents_dir.glob("*.docx.json")):
        spec = json.loads(spec_path.read_text(encoding="utf-8"))
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise RuntimeError("DOCX 评测 fixture 需要安装 python-docx") from exc
        with tempfile.TemporaryDirectory(prefix="rag-eval-docx-") as temp_dir:
            filename = spec_path.name.removesuffix(".json")
            docx_path = Path(temp_dir) / filename
            word = WordDocument()
            for block in spec.get("blocks", []):
                if block.get("type") == "heading":
                    word.add_heading(str(block.get("text") or ""), level=int(block.get("level") or 1))
                elif block.get("type") == "paragraph":
                    word.add_paragraph(str(block.get("text") or ""))
                elif block.get("type") == "table":
                    rows = block.get("rows") or []
                    if rows:
                        table = word.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                        for row_index, row in enumerate(rows):
                            for column_index, value in enumerate(row):
                                table.cell(row_index, column_index).text = str(value)
            word.save(docx_path)
            document = processor.parse_file(docx_path, original_name=filename)
            document.metadata["knowledge_base_id"] = str(spec.get("knowledge_base_id") or "default")
            retriever.add_document(document, processor.split(document))
    return RagEngine(retriever)


def _expected_sources(case: dict) -> set[str]:
    sources = case.get("expected_sources") or []
    if case.get("expected_source"):
        sources = [*sources, case["expected_source"]]
    return {str(source) for source in sources if source}


def _is_relevant(citation: dict, case: dict) -> bool:
    sources = _expected_sources(case)
    if sources and citation.get("filename") not in sources:
        return False
    keywords = [str(item).lower() for item in case.get("expected_keywords", []) if item]
    if not keywords:
        return bool(sources)
    haystack = f"{citation.get('text', '')}\n{citation.get('snippet', '')}".lower()
    return any(keyword in haystack for keyword in keywords)


def evaluate_cases(cases: list[dict], engine: RagEngine, top_k: int = 5) -> list[dict]:
    rows: list[dict] = []
    for case in cases:
        should_answer = bool(case.get("should_answer", True))
        context = [str(item) for item in case.get("conversation_context", []) if item]
        effective_question = "\n".join([*context[-4:], case["question"]])
        retrieval_options = {
            "query_rewrite": False,
            "knowledge_base_ids": case.get("knowledge_base_ids") or [],
            "min_score": case.get("min_score"),
            "strategy": case.get("strategy", "hybrid"),
            "graph_weight": case.get("graph_weight", 0.25),
            "graph_max_hops": case.get("graph_max_hops", 2),
        }
        retrieval = engine.search(
            effective_question,
            top_k=max(10, top_k),
            **retrieval_options,
        )
        result = engine.ask(
            effective_question,
            top_k=top_k,
            **retrieval_options,
        )
        citations = result.get("citations", [])
        retrieval_results = retrieval.get("results", [])
        trace = result.get("retrieval_trace", {})
        relevant_ranks = [
            index
            for index, citation in enumerate(retrieval_results, start=1)
            if _is_relevant(citation, case)
        ]
        refusal_reason = result.get("retrieval_trace", {}).get("refusal_reason")
        refused = bool(refusal_reason) or not citations
        citation_texts = [f"{item.get('text', '')}\n{item.get('snippet', '')}".lower() for item in citations]
        citation_audit = result.get("citation_audit") if isinstance(result.get("citation_audit"), dict) else {}
        valid_citations = [
            item for item in citations
            if item.get("id") and item.get("document_id") and str(item.get("text") or "").strip()
        ]
        expected_sources = _expected_sources(case)
        top_ten_sources = {str(item.get("filename") or "") for item in retrieval_results[:10]}
        expected_cells = [str(item).lower() for item in case.get("expected_cells", [])]
        expected_caption = [str(item).lower() for item in case.get("expected_caption_terms", [])]
        expected_formula = "".join(str(case.get("expected_formula") or "").lower().split())
        graph_payload = trace.get("pipeline", {}).get("graph", {})
        paths = graph_payload.get("paths") or []
        expected_path = [str(item).lower() for item in case.get("expected_path", [])]
        expected_relations = [str(item).lower() for item in case.get("expected_relations", [])]
        graph_path_correct = None
        graph_evidence_coverage = None
        if case.get("strategy") == "hybrid_graph":
            graph_path_correct = any(
                all(any(expected in str(label).lower() for label in path.get("labels", [])) for expected in expected_path)
                and all(expected in [str(item).lower() for item in path.get("relations", [])] for expected in expected_relations)
                for path in paths
            )
            graph_evidence = set(graph_payload.get("evidence_element_ids") or [])
            cited_evidence = {element_id for citation in citations for element_id in citation.get("element_ids", [])}
            graph_evidence_coverage = len(graph_evidence & cited_evidence) / max(len(graph_evidence), 1)
        rows.append(
            {
                "id": case["id"],
                "category": case.get("category", "general"),
                "question": case["question"],
                "should_answer": should_answer,
                "decision": "refused" if refused else "answered",
                "decision_correct": (not refused) if should_answer else refused,
                "recall_hit": any(rank <= top_k for rank in relevant_ranks) if should_answer else None,
                "first_relevant_rank": relevant_ranks[0] if relevant_ranks else None,
                "citation_correct": bool(citations and _is_relevant(citations[0], case)) if should_answer else not citations,
                "citation_coverage": float(citation_audit.get("coverage") or 0) if should_answer else None,
                "valid_citation_rate": len(valid_citations) / max(len(citations), 1) if citations else (1.0 if not should_answer else 0.0),
                "multihop_complete": expected_sources.issubset(top_ten_sources) if should_answer and case.get("category") == "multihop-graph" else None,
                "expected_sources": sorted(_expected_sources(case)),
                "top_sources": [citation.get("filename", "") for citation in citations[:3]],
                "refusal_reason": refusal_reason,
                "knowledge_base_ids": case.get("knowledge_base_ids") or [],
                "min_score": case.get("min_score"),
                "top_score": citations[0].get("rerank_score", 0) if citations else 0,
                "modality": case.get("modality"),
                "table_cell_correct": (
                    any(all(cell in text for cell in expected_cells) for text in citation_texts)
                    if expected_cells else None
                ),
                "caption_aligned": (
                    any(all(term in text for term in expected_caption) for text in citation_texts)
                    if expected_caption else None
                ),
                "formula_correct": (
                    any(expected_formula in "".join(text.split()) for text in citation_texts)
                    if expected_formula else None
                ),
                "graph_path_correct": graph_path_correct,
                "graph_evidence_coverage": graph_evidence_coverage,
            }
        )
    return rows


def summarize_rows(rows: list[dict], top_k: int = 5) -> dict:
    answerable = [row for row in rows if row["should_answer"]]
    refusal = [row for row in rows if not row["should_answer"]]
    def was_refused(row: dict) -> bool:
        if "decision" in row:
            return row.get("decision") == "refused"
        return bool(row.get("decision_correct")) if not row["should_answer"] else not bool(row.get("decision_correct"))
    recall = sum(bool(row["recall_hit"]) for row in answerable) / max(len(answerable), 1)
    mrr = sum(1 / row["first_relevant_rank"] if row["first_relevant_rank"] else 0 for row in answerable) / max(len(answerable), 1)
    citation_accuracy = sum(bool(row["citation_correct"]) for row in answerable) / max(len(answerable), 1)
    refusal_true_positive = sum(was_refused(row) for row in refusal)
    refusal_false_positive = sum(was_refused(row) for row in answerable)
    refusal_false_negative = sum(not was_refused(row) for row in refusal)
    refusal_precision = refusal_true_positive / max(refusal_true_positive + refusal_false_positive, 1)
    refusal_recall = refusal_true_positive / max(refusal_true_positive + refusal_false_negative, 1)
    refusal_f1 = 2 * refusal_precision * refusal_recall / max(refusal_precision + refusal_recall, 1e-12)
    misrefusal_rate = refusal_false_positive / max(len(answerable), 1)
    answer_acceptance = sum(bool(row["decision_correct"]) for row in answerable) / max(len(answerable), 1)
    modality = [row for row in answerable if row.get("modality")]
    table = [row for row in rows if row.get("table_cell_correct") is not None]
    captions = [row for row in rows if row.get("caption_aligned") is not None]
    formulas = [row for row in rows if row.get("formula_correct") is not None]
    graph = [row for row in rows if row.get("graph_path_correct") is not None]
    graph_coverage = [row for row in rows if row.get("graph_evidence_coverage") is not None]
    multihop = [row for row in answerable if row.get("multihop_complete") is not None]
    category_distribution = {
        category: sum(1 for row in rows if row.get("category", "general") == category)
        for category in sorted({row.get("category", "general") for row in rows})
    }
    source_distribution: dict[str, int] = {}
    for row in rows:
        for source in row.get("expected_sources", []):
            source_distribution[source] = source_distribution.get(source, 0) + 1
    return {
        "cases": len(rows),
        "answerable_cases": len(answerable),
        "refusal_cases": len(refusal),
        f"recall_at_{top_k}": round(recall, 4),
        "mrr": round(mrr, 4),
        "citation_accuracy": round(citation_accuracy, 4),
        "citation_coverage": round(sum(float(row.get("citation_coverage") or 0) for row in answerable) / max(len(answerable), 1), 4),
        "valid_citation_rate": round(sum(float(row.get("valid_citation_rate") or 0) for row in rows) / max(len(rows), 1), 4),
        "refusal_accuracy": round(refusal_recall, 4),
        "refusal_f1": round(refusal_f1, 4),
        "refusal_recall": round(refusal_recall, 4),
        "answerable_misrefusal_complement": round(1 - misrefusal_rate, 4),
        "answer_acceptance_accuracy": round(answer_acceptance, 4),
        "modality_recall_at_5": round(sum(bool(row["recall_hit"]) for row in modality) / max(len(modality), 1), 4),
        "modality_recall_at_10": round(sum(bool(row["first_relevant_rank"] and row["first_relevant_rank"] <= 10) for row in modality) / max(len(modality), 1), 4),
        "table_cell_accuracy": round(sum(bool(row["table_cell_correct"]) for row in table) / max(len(table), 1), 4),
        "caption_alignment": round(sum(bool(row["caption_aligned"]) for row in captions) / max(len(captions), 1), 4),
        "formula_accuracy": round(sum(bool(row["formula_correct"]) for row in formulas) / max(len(formulas), 1), 4),
        "graph_path_precision": round(sum(bool(row["graph_path_correct"]) for row in graph) / max(len(graph), 1), 4),
        "graph_evidence_coverage": round(sum(float(row["graph_evidence_coverage"]) for row in graph_coverage) / max(len(graph_coverage), 1), 4),
        "multihop_recall_at_5": round(sum(bool(row["recall_hit"]) for row in multihop) / max(len(multihop), 1), 4),
        "multihop_complete_evidence_at_10": round(sum(bool(row["multihop_complete"]) for row in multihop) / max(len(multihop), 1), 4),
        "category_distribution": category_distribution,
        "source_distribution": source_distribution,
    }


def compare_thresholds(summary: dict, thresholds: dict) -> list[dict]:
    checks = []
    for metric, minimum in thresholds.items():
        actual = float(summary.get(metric, 0))
        checks.append({"metric": metric, "actual": actual, "minimum": float(minimum), "passed": actual >= float(minimum)})
    return checks


def markdown_report(report: dict) -> str:
    lines = [
        "# 离线检索回归报告",
        "",
        f"生成时间：{report['generated_at']}",
        f"固定样本：{report['summary']['cases']}（可回答 {report['summary']['answerable_cases']} / 拒答 {report['summary']['refusal_cases']}）",
        "",
        "## 关键指标",
        "",
        "| 指标 | 实际 | 阈值 | 结果 |",
        "| --- | ---: | ---: | :---: |",
    ]
    for check in report["checks"]:
        label = METRIC_LABELS.get(check["metric"], check["metric"])
        lines.append(f"| {label} | {check['actual']:.4f} | {check['minimum']:.4f} | {'通过' if check['passed'] else '失败'} |")
    lines.extend([
        "",
        "## 分布",
        "",
        f"- 类别：{json.dumps(report['summary']['category_distribution'], ensure_ascii=False)}",
        f"- 期望来源：{json.dumps(report['summary']['source_distribution'], ensure_ascii=False)}",
    ])
    lines.extend(["", "## Case 明细", "", "| ID | 决策 | Recall | 首个相关排名 | 首条引用 | Top 来源 |", "| --- | --- | :---: | ---: | :---: | --- |"])
    for row in report["rows"]:
        recall = "—" if row["recall_hit"] is None else ("✓" if row["recall_hit"] else "✗")
        rank = row["first_relevant_rank"] or "—"
        sources = ", ".join(row["top_sources"]) or "—"
        lines.append(f"| {row['id']} | {row['decision']} | {recall} | {rank} | {'✓' if row['citation_correct'] else '✗'} | {sources} |")
    failed = [row for row in report["rows"] if not row["decision_correct"] or row["citation_correct"] is False or row["recall_hit"] is False]
    if failed:
        lines.extend(["", "## 需要处理", ""])
        for row in failed:
            lines.append(f"- `{row['id']}` {row['question']}（期望：{', '.join(row['expected_sources']) or '拒答'}；实际：{', '.join(row['top_sources']) or row['decision']}）")
    lines.append("")
    return "\n".join(lines)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="运行确定性的离线检索回归评测")
    parser.add_argument("--cases", type=Path, default=ROOT / "eval" / "cases.jsonl")
    parser.add_argument("--thresholds", type=Path, default=ROOT / "eval" / "thresholds.json")
    parser.add_argument("--documents", type=Path, default=ROOT / "samples" / "demo-documents")
    parser.add_argument("--report-dir", type=Path, default=ROOT / "eval" / "reports")
    parser.add_argument("--no-fail", action="store_true", help="写入基线报告，但不强制执行阈值")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    cases = load_jsonl(args.cases)
    thresholds = json.loads(args.thresholds.read_text(encoding="utf-8"))
    engine = build_offline_engine(args.documents)
    rows = evaluate_cases(cases, engine, top_k=5)
    summary = summarize_rows(rows, top_k=5)
    checks = compare_thresholds(summary, thresholds)
    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "mode": "offline-mock-template",
        "summary": summary,
        "thresholds": thresholds,
        "checks": checks,
        "passed": all(check["passed"] for check in checks),
        "rows": rows,
    }
    args.report_dir.mkdir(parents=True, exist_ok=True)
    json_path = args.report_dir / "latest.json"
    markdown_path = args.report_dir / "latest.md"
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    markdown_path.write_text(markdown_report(report), encoding="utf-8")
    print(json.dumps({**summary, "passed": report["passed"], "report": str(markdown_path)}, ensure_ascii=False, indent=2))
    return 0 if report["passed"] or args.no_fail else 1


if __name__ == "__main__":
    raise SystemExit(main())
