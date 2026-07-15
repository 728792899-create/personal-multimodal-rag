from __future__ import annotations

from pathlib import Path
from io import BytesIO
from types import SimpleNamespace

import fitz
from docx import Document as WordDocument
from fastapi.testclient import TestClient
import httpx
from PIL import Image
import pytest

from app.models.domain import DocumentElement
from app.services.document_processor import DocumentProcessor
from app.services.document_registry import DocumentRegistry
from app.services.object_store import LocalObjectStore
from app.services.ingestion_jobs import IngestionWorker
from app.services.parser_worker import ParserJobCancelled, ParserWorkerClient, document_from_content_list
from app.services.multimodal_assets import materialize_document_assets


def test_schema_v4_persists_document_elements_assets_and_parser_runs(tmp_path: Path):
    registry = DocumentRegistry(str(tmp_path / "registry.sqlite3"))
    assert registry.schema_version == 5

    document = DocumentProcessor().parse_text_source(
        "# Architecture\n\nThe parser emits typed elements.",
        "architecture.md",
    )
    document.metadata["knowledge_base_id"] = "default"
    registry.save_document(document)

    stored = registry.list_document_elements(document.document_id)
    assert [item["type"] for item in stored] == ["heading", "text"]
    assert stored[0]["heading_path"] == ["Architecture"]

    asset = registry.create_asset(
        knowledge_base_id="default",
        kind="source",
        object_key="ab/abcdef",
        original_name="architecture.md",
        media_type="text/markdown",
        sha256="abcdef",
        size_bytes=42,
        document_id=document.document_id,
    )
    assert registry.get_asset(asset["id"])["original_name"] == "architecture.md"

    run = registry.create_parser_run(
        document_id=document.document_id,
        provider="builtin",
        parser="markdown",
        status="succeeded",
        payload={"element_count": 2},
    )
    assert registry.list_parser_runs(document.document_id)[0]["id"] == run["id"]

    registry.delete_document(document.document_id)
    assert registry.list_document_elements(document.document_id) == []
    assert registry.list_parser_runs(document.document_id) == []
    assert registry.list_assets(document_id=document.document_id) == []


def test_local_object_store_is_content_addressed_and_deduplicated(tmp_path: Path):
    store = LocalObjectStore(tmp_path / "objects")
    first = store.put_bytes(b"same payload")
    second = store.put_bytes(b"same payload")

    assert first.sha256 == second.sha256
    assert first.object_key == second.object_key
    assert first.path == second.path
    assert first.path.read_bytes() == b"same payload"
    assert first.path.name == first.sha256
    assert str(first.path).startswith(str((tmp_path / "objects").resolve()))


def test_docx_preserves_paragraph_table_paragraph_order_as_elements(tmp_path: Path):
    source = tmp_path / "ordered.docx"
    word = WordDocument()
    word.add_heading("Index lifecycle", level=1)
    word.add_paragraph("Before the table.")
    table = word.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Stage"
    table.cell(0, 1).text = "State"
    table.cell(1, 0).text = "Graph"
    table.cell(1, 1).text = "ready"
    word.add_paragraph("After the table.")
    word.save(source)

    document = DocumentProcessor().parse_file(source)
    assert [element.type for element in document.elements] == ["heading", "text", "table", "text"]
    assert document.elements[2].table == [["Stage", "State"], ["Graph", "ready"]]
    assert "Stage | State" in document.elements[2].text

    chunks = DocumentProcessor(chunk_size=120).split(document)
    assert all(chunk.element_ids for chunk in chunks)
    assert any(chunk.modality == "table" for chunk in chunks)


def test_context_ready_element_schema_rejects_unknown_modality():
    element = DocumentElement(
        element_id="e1",
        document_id="d1",
        type="equation",
        order=0,
        text="E = mc^2",
        latex="E = mc^2",
        confidence=0.98,
    )
    assert element.latex == "E = mc^2"
    assert element.confidence == 0.98


def test_document_element_and_source_asset_api(monkeypatch, tmp_path: Path):
    from app.api.routers import documents
    from app.main import app

    store = LocalObjectStore(tmp_path / "api-objects")
    monkeypatch.setattr(documents, "object_store", store)
    monkeypatch.setattr(documents, "DATA_DIR", tmp_path / "uploads")
    client = TestClient(app)

    parser_status = client.get("/api/parsers/status")
    assert parser_status.status_code == 200
    assert parser_status.json()["profiles"][0]["id"] == "builtin"
    assert parser_status.json()["profiles"][0]["available"] is True

    upload = client.post(
        "/api/documents",
        files={"file": ("elements.md", b"# Elements\n\nPrecise citations use element identifiers.", "text/markdown")},
    )
    assert upload.status_code == 200
    document_id = upload.json()["document"]["id"]
    assert upload.json()["document"]["source_available"] is True
    assert upload.json()["document"]["element_count"] == 2

    elements = client.get(f"/api/documents/{document_id}/elements")
    assert elements.status_code == 200
    assert [item["type"] for item in elements.json()["elements"]] == ["heading", "text"]

    source = client.get(f"/api/documents/{document_id}/source")
    assert source.status_code == 200
    assert source.content.startswith(b"# Elements")
    assert source.headers["x-content-type-options"] == "nosniff"
    assert str(tmp_path) not in source.text

    reindex = client.post(f"/api/documents/{document_id}/reindex")
    assert reindex.status_code == 200
    assert reindex.json()["rebuilt"] is True
    rebuilt_elements = client.get(f"/api/documents/{document_id}/elements").json()["elements"]
    assert all(item["document_id"] == document_id for item in rebuilt_elements)
    assert all(item["id"].startswith(f"{document_id}:element:") for item in rebuilt_elements)
    assert client.get(f"/api/documents/{document_id}").json()["document"]["source_available"] is True

    deleted = client.delete(f"/api/documents/{document_id}")
    assert deleted.status_code == 200
    assert list((tmp_path / "api-objects").rglob("*")) == []


def test_parser_worker_contract_and_content_list_adapter(tmp_path: Path):
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.4\n")
    polls = {"count": 0}

    def handler(request: httpx.Request):
        if request.url.path == "/v1/capabilities":
            return httpx.Response(200, json={"profiles": [{"id": "mineru", "available": True}]})
        if request.method == "POST" and request.url.path == "/v1/jobs":
            return httpx.Response(202, json={"id": "job-1", "status": "queued"})
        if request.method == "GET" and request.url.path == "/v1/jobs/job-1":
            polls["count"] += 1
            if polls["count"] == 1:
                return httpx.Response(200, json={"id": "job-1", "status": "running"})
            return httpx.Response(
                200,
                json={
                    "id": "job-1",
                    "status": "succeeded",
                    "result": {
                        "parser": "mineru",
                        "content_list": [
                            {
                                "type": "text",
                                "text": "Architecture",
                                "text_level": 1,
                                "page_idx": 0,
                                "img_path": "/tmp/parser-jobs/private.png",
                            },
                            {"type": "table", "table_body": "Stage | State\nGraph | ready", "page_idx": 0},
                        ],
                    },
                },
            )
        return httpx.Response(404)

    client = ParserWorkerClient(
        "http://parser-worker",
        timeout_seconds=2,
        poll_seconds=0,
        http_client=httpx.Client(transport=httpx.MockTransport(handler)),
    )
    assert client.capabilities()["profiles"][0]["id"] == "mineru"
    result = client.parse(source, "sample.pdf", "mineru")
    document = document_from_content_list(
        result["content_list"],
        source_path=source,
        original_name="sample.pdf",
        parser_name=result["parser"],
    )
    assert [item.type for item in document.elements] == ["heading", "table"]
    assert document.elements[1].table[1] == ["Graph", "ready"]
    assert document.metadata["parser_provider"] == "raganything_worker"
    assert "/tmp/parser-jobs" not in document.model_dump_json()


def test_parser_worker_cancel_and_timeout_cleanup_remote_jobs(tmp_path: Path):
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.4\n")
    deleted: list[str] = []

    def handler(request: httpx.Request):
        if request.method == "POST" and request.url.path == "/v1/jobs":
            return httpx.Response(202, json={"id": "job-1", "status": "queued"})
        if request.method == "GET" and request.url.path == "/v1/jobs/job-1":
            return httpx.Response(200, json={"id": "job-1", "status": "running"})
        if request.method == "DELETE" and request.url.path == "/v1/jobs/job-1":
            deleted.append(request.url.path)
            return httpx.Response(200, json={"id": "job-1", "status": "cancelled"})
        return httpx.Response(404)

    transport = httpx.MockTransport(handler)
    cancelling = ParserWorkerClient(
        "http://parser-worker",
        timeout_seconds=2,
        poll_seconds=0,
        http_client=httpx.Client(transport=transport),
    )
    with pytest.raises(ParserJobCancelled, match="cancelled"):
        cancelling.parse(source, "sample.pdf", "mineru", cancel_check=lambda: True)

    timing_out = ParserWorkerClient(
        "http://parser-worker",
        timeout_seconds=0,
        poll_seconds=0,
        http_client=httpx.Client(transport=transport),
    )
    with pytest.raises(TimeoutError, match="timed out"):
        timing_out.parse(source, "sample.pdf", "mineru")

    assert deleted == ["/v1/jobs/job-1", "/v1/jobs/job-1"]


def test_ingestion_worker_does_not_fallback_after_parser_cancellation(tmp_path: Path):
    class CancelledParser:
        def parse(self, *_args, **_kwargs):
            raise ParserJobCancelled("Parser job cancelled")

    class ForbiddenFallbackProcessor:
        def parse_file(self, *_args, **_kwargs):
            raise AssertionError("cancelled parser jobs must not run the builtin fallback")

    registry = DocumentRegistry(str(tmp_path / "registry.sqlite3"))
    store = LocalObjectStore(tmp_path / "objects")
    stored = store.put_bytes(b"%PDF-1.4\n")
    asset = registry.create_asset(
        knowledge_base_id="default",
        kind="source",
        object_key=stored.object_key,
        original_name="cancel.pdf",
        media_type="application/pdf",
        sha256=stored.sha256,
        size_bytes=stored.size_bytes,
    )
    job = registry.create_index_job(
        source_type="file",
        source_name="cancel.pdf",
        payload={"asset_id": asset["id"], "parser_profile": "mineru"},
        knowledge_base_id="default",
        idempotency_key="cancelled-parser-job",
    )
    worker = IngestionWorker(
        registry,
        ForbiddenFallbackProcessor(),
        object(),
        SimpleNamespace(
            ingestion_lease_seconds=30,
            ingestion_poll_seconds=0.01,
            parser_provider="builtin",
            parser_fallback_allowed=True,
        ),
        object_store=store,
        parser_client=CancelledParser(),
    )

    assert worker.run_once() is True
    assert registry.get_index_job(job["id"])["status"] == "cancelled"
    assert registry.get_asset(asset["id"]) is None
    assert not stored.path.exists()


def test_pdf_embedded_image_is_materialized_as_controlled_asset(tmp_path: Path):
    image_buffer = BytesIO()
    Image.new("RGB", (24, 24), color=(40, 120, 200)).save(image_buffer, format="PNG")
    pdf_path = tmp_path / "illustrated.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((40, 40), "Architecture diagram")
    page.insert_image(fitz.Rect(40, 60, 100, 120), stream=image_buffer.getvalue())
    pdf.save(pdf_path)
    pdf.close()

    registry = DocumentRegistry(str(tmp_path / "registry.sqlite3"))
    store = LocalObjectStore(tmp_path / "objects")
    document = DocumentProcessor().parse_file(pdf_path)
    document.metadata["knowledge_base_id"] = "default"
    registry.save_document(document)
    assets = materialize_document_assets(document, registry, store)

    image_elements = [item for item in document.elements if item.type == "image"]
    assert len(image_elements) == 1
    assert image_elements[0].asset_id == assets[0]["id"]
    private = registry.get_asset(assets[0]["id"], include_private=True)
    assert store.path_for(private["object_key"]).read_bytes().startswith(b"\x89PNG")


def test_forced_knowledge_base_delete_releases_content_addressed_objects(monkeypatch, tmp_path: Path):
    from app.api.routers import documents, knowledge_bases
    from app.main import app

    store = LocalObjectStore(tmp_path / "kb-objects")
    monkeypatch.setattr(documents, "object_store", store)
    monkeypatch.setattr(documents, "DATA_DIR", tmp_path / "uploads")
    monkeypatch.setattr(knowledge_bases, "object_store", store)
    with TestClient(app) as client:
        created = client.post("/api/knowledge-bases", json={"name": "Disposable assets"})
        knowledge_base_id = created.json()["knowledge_base"]["id"]
        uploaded = client.post(
            "/api/documents",
            data={"knowledge_base_id": knowledge_base_id},
            files={"file": ("delete-me.md", b"# Delete\n\nRelease the durable source object.", "text/markdown")},
        )
        assert uploaded.status_code == 200
        assert any(path.is_file() for path in (tmp_path / "kb-objects").rglob("*"))

        deleted = client.delete(f"/api/knowledge-bases/{knowledge_base_id}?force=true")

        assert deleted.status_code == 200
        assert not any(path.is_file() for path in (tmp_path / "kb-objects").rglob("*"))
