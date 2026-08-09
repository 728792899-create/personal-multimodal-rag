from __future__ import annotations

import re
import uuid
import hashlib
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import fitz

from app.models.domain import Chunk, Document, DocumentElement, DocumentPage
from app.services.ocr import ImageOCRAdapter
from app.services.safe_logging import public_error_message


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".png", ".jpg", ".jpeg"}


class DocumentProcessor:
    def __init__(
        self,
        chunk_size: int | None = None,
        overlap: int = 80,
        min_chunk_size: int = 350,
        target_chunk_size: int = 650,
        max_chunk_size: int = 900,
        ocr_adapter: ImageOCRAdapter | None = None,
        docx_max_entries: int = 500,
        docx_max_uncompressed_bytes: int = 64 * 1024 * 1024,
        docx_max_compression_ratio: int = 200,
    ):
        # An explicit ``chunk_size`` keeps the legacy single-limit behaviour
        # used by tests and callers that deliberately request small chunks.
        if chunk_size is not None:
            selected = max(1, int(chunk_size))
            self.min_chunk_size = min(selected, max(1, selected // 2))
            self.target_chunk_size = selected
            self.max_chunk_size = selected
        else:
            self.min_chunk_size = max(1, int(min_chunk_size))
            self.target_chunk_size = max(
                self.min_chunk_size,
                int(target_chunk_size),
            )
            self.max_chunk_size = max(
                self.target_chunk_size,
                int(max_chunk_size),
            )
        self.chunk_size = self.target_chunk_size
        self.overlap = max(0, min(int(overlap), self.max_chunk_size - 1))
        self.ocr_adapter = ocr_adapter or ImageOCRAdapter()
        self.docx_max_entries = docx_max_entries
        self.docx_max_uncompressed_bytes = docx_max_uncompressed_bytes
        self.docx_max_compression_ratio = docx_max_compression_ratio

    def parse_file(self, file_path: Path, original_name: Optional[str] = None) -> Document:
        filename = original_name or file_path.name
        # Content-addressed objects intentionally have no user-controlled
        # extension. Parser selection therefore follows the already validated
        # display name when one is provided.
        suffix = Path(filename).suffix.lower()
        document_id = str(uuid.uuid4())
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型：{suffix}")

        if suffix == ".pdf":
            pages, metadata, elements = self._parse_pdf(file_path, document_id)
        elif suffix == ".docx":
            pages, metadata, elements = self._parse_docx(file_path, document_id)
        elif suffix in {".png", ".jpg", ".jpeg"}:
            pages, metadata, elements = self._parse_image(file_path, document_id, filename)
        else:
            pages, metadata, elements = self._parse_text_file(file_path, suffix, document_id)

        if not any(page.text.strip() for page in pages):
            raise ValueError("未能从文档中提取可读文本。")

        metadata.update(
            {
                "file_size": file_path.stat().st_size,
                "content_hash": self._content_hash(file_path),
                "parser": metadata.get("parser", "plain_text"),
                "index_status": "parsed",
            }
        )

        return Document(
            document_id=document_id,
            file_name=filename,
            file_path=str(file_path),
            file_type=self._file_type(suffix),
            title=Path(filename).stem,
            created_at=datetime.utcnow(),
            pages=pages,
            elements=elements,
            metadata=metadata,
        )

    def parse_text_source(
        self,
        text: str,
        source_name: str,
        source_url: str = "",
        parser: str = "plain_text",
        metadata: Optional[dict] = None,
    ) -> Document:
        normalized = self.normalize_text(text)
        if not normalized:
            raise ValueError("未能从文档中提取可读文本。")
        payload = {
            "parser": parser,
            "index_status": "parsed",
            "source_url": source_url,
            "content_hash": hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
            **(metadata or {}),
        }
        document_id = str(uuid.uuid4())
        page = DocumentPage(page_number=None, text=normalized, metadata={"source_url": source_url})
        return Document(
            document_id=document_id,
            file_name=source_name,
            file_path=source_url or source_name,
            file_type="text",
            title=Path(source_name).stem,
            created_at=datetime.utcnow(),
            pages=[page],
            elements=self._text_elements(document_id, normalized, None),
            metadata=payload,
        )

    def split(self, doc: Document) -> list[Chunk]:
        paragraphs = self._paragraphs(doc)
        chunks: list[Chunk] = []
        current = ""
        current_meta = {"page_number": None, "heading_path": [], "element_ids": [], "modality": "text", "metadata": {}}
        current_boundary: tuple[int | None, tuple[str, ...]] | None = None

        for paragraph, meta in paragraphs:
            if meta.get("modality") in {"image", "table", "equation", "code"}:
                if current.strip():
                    chunks.append(self._make_chunk(doc, len(chunks), current, current_meta))
                    current = ""
                    current_boundary = None
                parts = (
                    self._split_table(paragraph, meta)
                    if meta.get("modality") == "table"
                    else [
                        (part, dict(meta))
                        for part in (
                            self._hard_split(paragraph)
                            if len(paragraph) > self.max_chunk_size
                            else [paragraph]
                        )
                    ]
                )
                for part, part_meta in parts:
                    chunks.append(
                        self._make_chunk(doc, len(chunks), part, part_meta)
                    )
                current_meta = {"page_number": None, "heading_path": [], "element_ids": [], "modality": "text", "metadata": {}}
                continue

            boundary = (
                meta.get("page_number"),
                tuple(str(item) for item in meta.get("heading_path", [])),
            )
            if current.strip() and boundary != current_boundary:
                chunks.append(self._make_chunk(doc, len(chunks), current, current_meta))
                current = ""
                current_boundary = None

            if len(paragraph) > self.max_chunk_size:
                if current.strip():
                    chunks.append(self._make_chunk(doc, len(chunks), current, current_meta))
                    current = ""
                for part in self._hard_split(paragraph):
                    chunks.append(self._make_chunk(doc, len(chunks), part, meta))
                current_boundary = None
                continue

            had_current = bool(current)
            candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
            if (
                len(candidate) <= self.target_chunk_size
                or (
                    current
                    and len(current) < self.min_chunk_size
                    and len(candidate) <= self.max_chunk_size
                )
            ):
                current = candidate
                current_meta = (
                    self._merge_chunk_meta(current_meta, meta)
                    if had_current
                    else dict(meta)
                )
                current_boundary = boundary
            else:
                if current.strip():
                    chunks.append(self._make_chunk(doc, len(chunks), current, current_meta))
                carry = self._tail(current) if current_boundary == boundary else ""
                next_text = f"{carry}\n\n{paragraph}".strip() if carry else paragraph
                if len(next_text) > self.max_chunk_size:
                    carry = ""
                    next_text = paragraph
                current = next_text
                current_meta = (
                    self._merge_chunk_meta(current_meta, meta)
                    if carry
                    else dict(meta)
                )
                current_boundary = boundary

        if current.strip():
            chunks.append(self._make_chunk(doc, len(chunks), current, current_meta))

        return chunks

    def normalize_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _content_hash(self, file_path: Path) -> str:
        digest = hashlib.sha256()
        with file_path.open("rb") as handle:
            for block in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(block)
        return digest.hexdigest()

    def _parse_pdf(self, file_path: Path, document_id: str) -> tuple[list[DocumentPage], dict, list[DocumentElement]]:
        pages: list[DocumentPage] = []
        elements: list[DocumentElement] = []
        try:
            pdf = fitz.open(str(file_path))
            for index, page in enumerate(pdf, start=1):
                page_parts: list[str] = []
                blocks = sorted(page.get_text("blocks"), key=lambda block: (round(block[1], 2), round(block[0], 2)))
                for block in blocks:
                    text = self.normalize_text(str(block[4] or ""))
                    if not text:
                        continue
                    page_parts.append(text)
                    elements.append(
                        DocumentElement(
                            element_id=f"{document_id}:element:{len(elements)}",
                            document_id=document_id,
                            type="text",
                            order=len(elements),
                            text=text,
                            page_number=index,
                            bbox=[round(float(value), 2) for value in block[:4]],
                            confidence=1.0,
                            metadata={"parser": "pymupdf", "block_type": int(block[6])},
                        )
                    )
                for image_index, image in enumerate(page.get_images(full=True)):
                    elements.append(
                        DocumentElement(
                            element_id=f"{document_id}:element:{len(elements)}",
                            document_id=document_id,
                            type="image",
                            order=len(elements),
                            text=f"第 {index} 页的内嵌图片",
                            page_number=index,
                            confidence=None,
                            metadata={"xref": int(image[0]), "page_image_index": image_index, "asset_status": "pending_materialization"},
                        )
                    )
                text = "\n\n".join(page_parts)
                if text:
                    pages.append(
                        DocumentPage(
                            page_number=index,
                            text=text,
                            metadata={"page_index": index - 1},
                        )
                    )
            metadata = {"page_count": pdf.page_count, "parser": "pymupdf"}
            pdf.close()
            return pages, metadata, elements
        except Exception as exc:
            raise ValueError(f"PDF 解析失败（{type(exc).__name__}）。") from exc

    def _parse_text_file(self, file_path: Path, suffix: str, document_id: str) -> tuple[list[DocumentPage], dict, list[DocumentElement]]:
        text = self.normalize_text(file_path.read_text(encoding="utf-8", errors="ignore"))
        metadata = {"parser": "markdown" if suffix in {".md", ".markdown"} else "plain_text"}
        return [DocumentPage(page_number=None, text=text, metadata={})], metadata, self._text_elements(document_id, text, None)

    def _parse_docx(self, file_path: Path, document_id: str) -> tuple[list[DocumentPage], dict, list[DocumentElement]]:
        self._validate_docx_archive(file_path)
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise ValueError("解析 DOCX 需要安装 python-docx。") from exc

        try:
            word = WordDocument(str(file_path))
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            from docx.oxml.ns import qn

            blocks: list[str] = []
            elements: list[DocumentElement] = []
            heading_path: list[str] = []
            heading_count = 0
            table_count = 0

            for child in word.element.body.iterchildren():
                if child.tag == qn("w:p"):
                    paragraph = Paragraph(child, word)
                    text = self.normalize_text(paragraph.text)
                    math_text = " ".join(
                        self.normalize_text(str(node.text or ""))
                        for node in child.iter()
                        if node.tag.endswith("}t") and "officeDocument/2006/math" in node.tag and node.text
                    ).strip()
                    if text:
                        style_name = str(getattr(paragraph.style, "name", "") or "")
                        match = re.match(r"Heading\s+([1-6])", style_name, flags=re.IGNORECASE)
                        element_type = "text"
                        if match:
                            level = int(match.group(1))
                            heading_count += 1
                            heading_path = heading_path[: level - 1] + [text]
                            rendered = f"{'#' * level} {text}"
                            element_type = "heading"
                        else:
                            rendered = text
                        blocks.append(rendered)
                        elements.append(
                            DocumentElement(
                                element_id=f"{document_id}:element:{len(elements)}",
                                document_id=document_id,
                                type=element_type,
                                order=len(elements),
                                text=rendered,
                                heading_path=list(heading_path),
                                confidence=1.0,
                                metadata={"style": style_name},
                            )
                        )
                    if math_text:
                        blocks.append(math_text)
                        elements.append(
                            DocumentElement(
                                element_id=f"{document_id}:element:{len(elements)}",
                                document_id=document_id,
                                type="equation",
                                order=len(elements),
                                text=math_text,
                                latex=math_text,
                                heading_path=list(heading_path),
                                confidence=0.75,
                                metadata={"source": "omml_text"},
                            )
                        )
                    relationship_ids = [
                        node.get(qn("r:embed"))
                        for node in child.iter()
                        if node.tag == qn("a:blip") and node.get(qn("r:embed"))
                    ]
                    for relationship_id in relationship_ids:
                        rendered = f"段落 {len(elements) + 1} 附近的内嵌图片"
                        blocks.append(rendered)
                        elements.append(
                            DocumentElement(
                                element_id=f"{document_id}:element:{len(elements)}",
                                document_id=document_id,
                                type="image",
                                order=len(elements),
                                text=rendered,
                                heading_path=list(heading_path),
                                metadata={"relationship_id": relationship_id, "asset_status": "pending_materialization"},
                            )
                        )
                elif child.tag == qn("w:tbl"):
                    table = Table(child, word)
                    structured: list[list[str]] = []
                    for row in table.rows:
                        cells = [self.normalize_text(cell.text).replace("\n", " ") for cell in row.cells]
                        if any(cells):
                            structured.append(cells)
                    if structured:
                        table_count += 1
                        rendered = "\n".join(" | ".join(row) for row in structured)
                        blocks.append(rendered)
                        elements.append(
                            DocumentElement(
                                element_id=f"{document_id}:element:{len(elements)}",
                                document_id=document_id,
                                type="table",
                                order=len(elements),
                                text=rendered,
                                table=structured,
                                heading_path=list(heading_path),
                                confidence=1.0,
                                metadata={"row_count": len(structured), "column_count": max(len(row) for row in structured)},
                            )
                        )
            text = "\n\n".join(blocks)
            return [DocumentPage(page_number=None, text=text, metadata={"table_count": table_count})], {
                "parser": "python-docx",
                "heading_count": heading_count,
                "table_count": table_count,
            }, elements
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"DOCX 解析失败（{type(exc).__name__}）。") from exc

    def _validate_docx_archive(self, file_path: Path) -> None:
        if not zipfile.is_zipfile(file_path):
            raise ValueError("DOCX 不是有效的 Office 文档压缩包。")
        try:
            with zipfile.ZipFile(file_path) as archive:
                entries = archive.infolist()
                names = {entry.filename for entry in entries}
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise ValueError("DOCX 不是有效的 Office 文档压缩包。")
                if len(entries) > self.docx_max_entries:
                    raise ValueError("DOCX 压缩包包含过多条目。")
                expanded_size = sum(max(0, entry.file_size) for entry in entries)
                if expanded_size > self.docx_max_uncompressed_bytes:
                    raise ValueError("DOCX 解压后大小超过配置上限。")
                for entry in entries:
                    if entry.file_size <= 4096:
                        continue
                    ratio = entry.file_size / max(entry.compress_size, 1)
                    if ratio > self.docx_max_compression_ratio:
                        raise ValueError("DOCX 压缩比异常，已拒绝解析。")
        except zipfile.BadZipFile as exc:
            raise ValueError("DOCX 不是有效的 Office 文档压缩包。") from exc

    def _parse_image(
        self,
        file_path: Path,
        document_id: str,
        display_name: str,
    ) -> tuple[list[DocumentPage], dict, list[DocumentElement]]:
        from PIL import Image

        with Image.open(file_path) as image:
            image_format = str(image.format or "").lower()
            width, height = image.size
        result = self.ocr_adapter.extract_text(file_path)
        if result.text:
            text = f"图片文件：{display_name}\n\nOCR 提取文本：\n{self.normalize_text(result.text)}"
        else:
            text = (
                f"图片文件：{display_name}。\n"
                "当前环境未完成 OCR 文本提取，已记录图片元数据；安装 tesseract + pytesseract 后可自动提取图片文本。"
            )
        public_error = (
            public_error_message(
                result.error,
                "OCR 运行环境不可用。"
                if result.status == "unavailable"
                else "OCR 文本提取失败。",
            )
            if result.error
            else ""
        )
        page = DocumentPage(
            page_number=None,
            text=text,
            metadata={
                "image_format": image_format,
                "image_width": width,
                "image_height": height,
                "ocr_status": result.status,
                "ocr_engine": result.engine,
                "ocr_error": public_error,
            },
        )
        element = DocumentElement(
            element_id=f"{document_id}:element:0",
            document_id=document_id,
            type="image",
            order=0,
            text=text,
            confidence=1.0 if result.text else 0.0,
            metadata={**page.metadata, "source_asset_role": True},
        )
        return [page], {"parser": "image_ocr", "ocr_status": result.status, "ocr_engine": result.engine}, [element]

    def _paragraphs(self, doc: Document) -> list[tuple[str, dict]]:
        paragraphs: list[tuple[str, dict]] = []
        if doc.elements:
            for element in sorted(doc.elements, key=lambda item: item.order):
                if not element.text.strip():
                    continue
                # Headings live in ``heading_path`` and embedding text. They do
                # not become citation-visible leaf evidence on their own.
                if element.type == "heading":
                    continue
                retrieval_text = element.text.strip()
                enrichment = element.metadata.get("enrichment")
                if isinstance(enrichment, dict):
                    description = str(enrichment.get("description") or "").strip()
                    keywords = [str(item) for item in enrichment.get("keywords", []) if str(item).strip()]
                    if description and description not in retrieval_text:
                        retrieval_text = f"{retrieval_text}\n\n{description}"
                    if keywords:
                        retrieval_text = f"{retrieval_text}\n\nKeywords: {', '.join(keywords[:16])}"
                paragraphs.append(
                    (
                        retrieval_text,
                        {
                            "page_number": element.page_number,
                            "heading_path": list(element.heading_path),
                            "element_ids": [element.element_id],
                            "modality": element.type,
                            "metadata": {
                                "element_ids": [element.element_id],
                                "modality": element.type,
                                "bbox": element.bbox,
                                "asset_id": element.asset_id,
                                "table_rows": element.table,
                            },
                        },
                    )
                )
            return paragraphs
        heading_path: list[str] = []
        for page in doc.pages:
            for part in re.split(r"\n\s*\n", page.text):
                paragraph = part.strip()
                if not paragraph:
                    continue
                heading = re.match(r"^(#{1,6})\s+(.+)$", paragraph)
                if heading:
                    level = len(heading.group(1))
                    heading_path = heading_path[: level - 1] + [heading.group(2).strip()]
                paragraphs.append(
                    (
                        paragraph,
                        {
                            "page_number": page.page_number,
                            "heading_path": list(heading_path),
                            "element_ids": [],
                            "modality": "text",
                            "metadata": dict(page.metadata),
                        },
                    )
                )
        return paragraphs

    def _hard_split(self, text: str) -> list[str]:
        parts: list[str] = []
        start = 0
        while start < len(text):
            hard_end = min(len(text), start + self.max_chunk_size)
            target_end = min(len(text), start + self.target_chunk_size)
            minimum_end = min(len(text), start + self.min_chunk_size)
            end = self._preferred_boundary(
                text,
                minimum_end=minimum_end,
                target_end=target_end,
                hard_end=hard_end,
            )
            if end <= start:
                end = hard_end
            parts.append(text[start:end].strip())
            if end >= len(text):
                break
            start = max(end - self.overlap, start + 1)
        return [part for part in parts if part]

    @staticmethod
    def _preferred_boundary(
        text: str,
        *,
        minimum_end: int,
        target_end: int,
        hard_end: int,
    ) -> int:
        boundaries = [
            match.end()
            for match in re.finditer(r"[\n。！？；.!?;]", text[minimum_end:hard_end])
        ]
        absolute = [minimum_end + offset for offset in boundaries]
        before_target = [position for position in absolute if position <= target_end]
        if before_target:
            return before_target[-1]
        if absolute:
            return absolute[0]
        return hard_end

    def _split_table(self, text: str, meta: dict) -> list[tuple[str, dict]]:
        if len(text) <= self.max_chunk_size:
            return [(text, dict(meta))]
        raw_rows = meta.get("metadata", {}).get("table_rows")
        rows = [
            " | ".join(str(cell) for cell in row)
            for row in raw_rows
            if isinstance(row, list)
        ] if isinstance(raw_rows, list) else []
        if len(rows) < 2:
            rows = [row for row in text.splitlines() if row.strip()]
        if len(rows) < 2:
            return [(part, dict(meta)) for part in self._hard_split(text)]

        header = rows[0]
        windows: list[tuple[str, dict]] = []
        current_rows: list[str] = []
        row_start = 1
        for row_index, row in enumerate(rows[1:], start=1):
            candidate = "\n".join([header, *current_rows, row])
            if current_rows and len(candidate) > self.target_chunk_size:
                windows.append(
                    self._table_window(
                        header,
                        current_rows,
                        meta,
                        row_start=row_start,
                        row_end=row_index - 1,
                    )
                )
                current_rows = []
                row_start = row_index
            current_rows.append(row)
        if current_rows:
            windows.append(
                self._table_window(
                    header,
                    current_rows,
                    meta,
                    row_start=row_start,
                    row_end=len(rows) - 1,
                )
            )
        expanded: list[tuple[str, dict]] = []
        for rendered, window_meta in windows:
            if len(rendered) <= self.max_chunk_size:
                expanded.append((rendered, window_meta))
            else:
                expanded.extend(
                    (part, dict(window_meta)) for part in self._hard_split(rendered)
                )
        return expanded

    @staticmethod
    def _table_window(
        header: str,
        rows: list[str],
        meta: dict,
        *,
        row_start: int,
        row_end: int,
    ) -> tuple[str, dict]:
        window_meta = {
            **meta,
            "metadata": {
                **meta.get("metadata", {}),
                "table_row_start": row_start,
                "table_row_end": row_end,
                "table_header_repeated": True,
            },
        }
        return "\n".join([header, *rows]), window_meta

    def _tail(self, text: str) -> str:
        if not text:
            return ""
        return text[-self.overlap :]

    def _make_chunk(self, doc: Document, index: int, text: str, meta: dict) -> Chunk:
        normalized = text.strip()
        heading_path = [str(item) for item in meta.get("heading_path", [])]
        embedding_parts = [str(doc.title or doc.file_name).strip()]
        if heading_path:
            embedding_parts.append(" > ".join(heading_path))
        embedding_parts.append(normalized)
        source_metadata = dict(meta.get("metadata", {}))
        source_metadata.pop("table_rows", None)
        metadata = {
            **source_metadata,
            "embedding_text": "\n\n".join(
                part for part in embedding_parts if part
            ),
            "knowledge_base_id": str(
                doc.metadata.get("knowledge_base_id") or "default"
            ),
            "content_hash": str(doc.metadata.get("content_hash") or ""),
            "chunker_version": "structure-v2",
            "parser_version": str(
                doc.metadata.get("parser_version")
                or doc.metadata.get("parser")
                or "builtin"
            ),
        }
        return Chunk(
            chunk_id=f"{doc.document_id}:{index}",
            document_id=doc.document_id,
            file_name=doc.file_name,
            chunk_index=index,
            text=normalized,
            page_number=meta.get("page_number"),
            heading_path=heading_path,
            element_ids=meta.get("element_ids", []),
            modality=meta.get("modality", "text"),
            parent_element_id=(meta.get("element_ids") or [None])[0],
            metadata=metadata,
        )

    def _text_elements(self, document_id: str, text: str, page_number: int | None) -> list[DocumentElement]:
        elements: list[DocumentElement] = []
        heading_path: list[str] = []
        for part in re.split(r"\n\s*\n", text):
            normalized = part.strip()
            if not normalized:
                continue
            match = re.match(r"^(#{1,6})\s+(.+)$", normalized)
            element_type = "text"
            if match:
                level = len(match.group(1))
                heading_path = heading_path[: level - 1] + [match.group(2).strip()]
                element_type = "heading"
            elements.append(
                DocumentElement(
                    element_id=f"{document_id}:element:{len(elements)}",
                    document_id=document_id,
                    type=element_type,
                    order=len(elements),
                    text=normalized,
                    page_number=page_number,
                    heading_path=list(heading_path),
                    confidence=1.0,
                )
            )
        return elements

    @staticmethod
    def _merge_chunk_meta(current: dict, next_meta: dict) -> dict:
        if not current.get("element_ids"):
            return dict(next_meta)
        element_ids = list(dict.fromkeys([*(current.get("element_ids") or []), *(next_meta.get("element_ids") or [])]))
        modalities = {current.get("modality", "text"), next_meta.get("modality", "text")}
        metadata = {**current.get("metadata", {}), **next_meta.get("metadata", {}), "element_ids": element_ids}
        return {
            "page_number": current.get("page_number") if current.get("page_number") is not None else next_meta.get("page_number"),
            "heading_path": next_meta.get("heading_path") or current.get("heading_path", []),
            "element_ids": element_ids,
            "modality": modalities.pop() if len(modalities) == 1 else "mixed",
            "metadata": metadata,
        }

    def _file_type(self, suffix: str):
        if suffix == ".pdf":
            return "pdf"
        if suffix == ".docx":
            return "docx"
        if suffix in {".md", ".markdown"}:
            return "markdown"
        if suffix in {".png", ".jpg", ".jpeg"}:
            return "image"
        return "text"
